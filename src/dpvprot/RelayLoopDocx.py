# Copyright (C) 2018-2021 Battelle Memorial Institute
# file: RelayLoopDocx.py
""" Insert PNGs and table for analysis of T400L from ATP COMTRADE files.

Paragraph.

Public Functions:
    :main: does the work
"""

import glob
import csv
from docx import Document
from docx.shared import Inches

def cellval(tok):
    if float(tok) > 0:
        return tok
    return 'n/a'

document = Document()
document.sections[-1].left_margin = Inches(1)
document.sections[-1].right_margin = Inches(1)
document.sections[-1].top_margin = Inches(0.75)
document.sections[-1].bottom_margin = Inches(0.75)

table = document.add_table(rows=1, cols=10)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Case'
hdr_cells[1].text = 'Site'
hdr_cells[2].text = 'Tfault'
hdr_cells[3].text = 'T45'
hdr_cells[4].text = 'T60'
hdr_cells[5].text = 'T88'
hdr_cells[6].text = 'Q46'
hdr_cells[7].text = 'LEN'
hdr_cells[8].text = 'Q47'
hdr_cells[9].text = 'LEN'
#hdr_cells[10].text = 'TD21P'
#hdr_cells[11].text = 'LEN'
#hdr_cells[12].text = 'TD21G'
#hdr_cells[13].text = 'LEN'
with open('RelayTimesSHE.csv', mode='r') as infile:
    reader = csv.reader(infile)
    for row in reader:
        png_name = row[0][:-4]
        site_name = row[1]
        tfault = row[2]
        row_cells = table.add_row().cells
        row_cells[0].text = png_name
        row_cells[1].text = site_name
        row_cells[2].text = tfault
        for i in range (3, 10):
            row_cells[i].text = cellval(row[i])
document.add_page_break()

files = sorted(glob.glob ('*.png'))
fignum = 1
for fname in files:
    document.add_picture(fname, width=Inches(6.5))
    document.add_paragraph('Figure ' + str(fignum) + ': ' + fname, style='Caption')
    document.add_page_break()
    fignum += 1

document.save('relay_plots_she.docx')
