# Copyright (C) 2018-2021 Battelle Memorial Institute
# file: ComtradeLoopDocx.py
""" Collects PNGs from ATP-generated COMTRADE files into a Word document.

Paragraph.

Public Functions:
    :main: does the work
"""

import operator
import subprocess
import os
import shutil
import glob
from docx import Document
from docx.shared import Inches

document = Document()
document.sections[-1].left_margin = Inches(1)
document.sections[-1].right_margin = Inches(1)

files = sorted(glob.glob ('J*.png'))
fignum = 1
for fname in files:
	document.add_picture(fname, width=Inches(6.5))
	document.add_paragraph('Figure ' + str(fignum) + ': ' + fname, style='Caption')
	document.add_page_break()
	fignum += 1

document.save('J1_case_plots_1MHz.docx')
