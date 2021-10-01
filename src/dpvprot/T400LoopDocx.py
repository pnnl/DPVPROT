# Copyright (C) 2018-2021 Battelle Memorial Institute
# file: T400LoopDocx.py
""" Assembles PNG plots and tabulated CSV file values into a document.

Paragraph.

Public Functions:
    :main: does the work
"""

import glob
import csv
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT

def cellval(tok):
    if float(tok) > 0:
        return tok
    return 'n/a'

document = Document()
for section in document.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

table = document.add_table(rows=1, cols=8)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Case'
hdr_cells[1].text = 'Site'
hdr_cells[2].text = 'Tfault'
hdr_cells[3].text = 'TD32F'
hdr_cells[4].text = 'OC21P'
hdr_cells[5].text = 'OC21G'
hdr_cells[6].text = 'TD21P'
hdr_cells[7].text = 'TD21G'
with open('T400LTimes.csv', mode='r') as infile:
    reader = csv.reader(infile)
    for row in reader:
        png_name = row[0]
        site_name = row[1]
        tfault = row[2]
        row_cells = table.add_row().cells
        row_cells[0].text = png_name
        row_cells[1].text = site_name
        row_cells[2].text = tfault
        for i in range (3, 8):
            row_cells[i].text = cellval(row[i])
document.add_page_break()

files = sorted(glob.glob ('*.png'))
fignum = 1
for fname in files:
    document.add_picture(fname, width=Inches(9.5))
    document.add_paragraph('Figure ' + str(fignum) + ': ' + fname, style='Caption')
    document.add_page_break()
    fignum += 1

document.save('T400L_plots_new.docx')
